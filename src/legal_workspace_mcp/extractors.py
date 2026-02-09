"""Text extraction from various document formats."""

import csv
import io
import json
import logging
import re
from pathlib import Path

import chardet

logger = logging.getLogger(__name__)


def extract_text(file_path: Path) -> str | None:
    """Extract text content from a file based on its extension.

    Args:
        file_path: Path to the file to extract text from.

    Returns:
        Extracted text content, or None if extraction fails.
    """
    suffix = file_path.suffix.lower()
    try:
        if suffix in (".txt", ".md", ".markdown", ".rtf"):
            return _extract_plain_text(file_path)
        elif suffix == ".pdf":
            return _extract_pdf(file_path)
        elif suffix == ".docx":
            return _extract_docx(file_path)
        elif suffix in (".html", ".htm"):
            return _extract_html(file_path)
        elif suffix == ".json":
            return _extract_json(file_path)
        elif suffix == ".csv":
            return _extract_csv(file_path)
        else:
            logger.warning("Unsupported file extension: %s for %s", suffix, file_path)
            return None
    except Exception as e:
        logger.error("Failed to extract text from %s: %s", file_path, e)
        return None


def _read_with_encoding(file_path: Path) -> str:
    """Read a file, detecting encoding if needed."""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    except UnicodeDecodeError:
        raw = file_path.read_bytes()
        detected = chardet.detect(raw)
        encoding = detected.get("encoding", "utf-8") or "utf-8"
        return raw.decode(encoding, errors="replace")


def _extract_plain_text(file_path: Path) -> str:
    """Extract text from plain text, markdown, or RTF files."""
    return _read_with_encoding(file_path)


def _extract_pdf(file_path: Path) -> str:
    """Extract text from PDF files using PyMuPDF (fitz)."""
    import fitz  # PyMuPDF

    text_parts: list[str] = []
    with fitz.open(str(file_path)) as doc:
        for page_num, page in enumerate(doc, 1):
            page_text = page.get_text("text")
            if page_text.strip():
                text_parts.append(f"[Page {page_num}]\n{page_text}")
    return "\n\n".join(text_parts)


def _extract_docx(file_path: Path) -> str:
    """Extract text from Word .docx files."""
    from docx import Document

    doc = Document(str(file_path))
    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Preserve heading structure
            if para.style and para.style.name and para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "").strip()
                try:
                    hashes = "#" * int(level)
                    parts.append(f"{hashes} {text}")
                except ValueError:
                    parts.append(text)
            else:
                parts.append(text)

    # Also extract text from tables
    for table in doc.tables:
        table_rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_rows.append(" | ".join(cells))
        if table_rows:
            parts.append("\n".join(table_rows))

    return "\n\n".join(parts)


def _extract_html(file_path: Path) -> str:
    """Extract text from HTML files."""
    from bs4 import BeautifulSoup

    content = _read_with_encoding(file_path)
    soup = BeautifulSoup(content, "html.parser")

    # Remove script and style elements
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()

    text = soup.get_text(separator="\n")
    # Collapse multiple blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _extract_json(file_path: Path) -> str:
    """Extract text from JSON files by pretty-printing."""
    content = _read_with_encoding(file_path)
    try:
        data = json.loads(content)
        return json.dumps(data, indent=2, ensure_ascii=False)
    except json.JSONDecodeError:
        return content


def _extract_csv(file_path: Path) -> str:
    """Extract text from CSV files, formatting as a readable table."""
    content = _read_with_encoding(file_path)
    reader = csv.reader(io.StringIO(content))
    rows = list(reader)
    if not rows:
        return ""

    # Format as markdown-like table
    parts: list[str] = []
    header = rows[0]
    parts.append(" | ".join(header))
    parts.append(" | ".join(["---"] * len(header)))
    for row in rows[1:]:
        parts.append(" | ".join(row))
    return "\n".join(parts)
